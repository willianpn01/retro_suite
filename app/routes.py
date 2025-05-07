from flask import Blueprint, render_template, request, send_file, jsonify
import os
import os
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from bs4 import BeautifulSoup
import base64
import requests
from PIL import Image
import re

bp = Blueprint('main', __name__)

@bp.route('/import_docx', methods=['POST'])
def import_docx():
    file = request.files.get('file')
    if not file:
        return jsonify({'error': 'Arquivo não enviado'}), 400
    from docx import Document as DocxDocument
    import html, base64
    from io import BytesIO
    doc = DocxDocument(file)
    html_out = []
    # Detectar listas
    in_ul = False
    in_ol = False
    for para in doc.paragraphs:
        style = para.style.name.lower()
        txt = ''
        for run in para.runs:
            t = html.escape(run.text)
            if not t:
                continue
            # Aplica negrito, itálico, sublinhado
            if run.bold:
                t = f'<b>{t}</b>'
            if run.italic:
                t = f'<i>{t}</i>'
            if run.underline:
                t = f'<u>{t}</u>'
            # Aplica tamanho da fonte
            font_size = None
            try:
                if run.font.size:
                    # run.font.size é em EMUs (1 pt = 12700 EMUs)
                    pt = run.font.size.pt
                    if pt:
                        px = int(pt * 1.33)  # 1pt ≈ 1.33px
                        t = f'<span style="font-size:{px}px">{t}</span>'
            except Exception:
                pass
            txt += t
        # Listas numeradas
        if 'list' in style or 'bullet' in style or 'num' in style:
            if 'bullet' in style or 'list' in style:
                if not in_ul:
                    html_out.append('<ul>')
                    in_ul = True
                html_out.append(f'<li>{txt}</li>')
            elif 'num' in style:
                if not in_ol:
                    html_out.append('<ol>')
                    in_ol = True
                html_out.append(f'<li>{txt}</li>')
        else:
            if in_ul:
                html_out.append('</ul>')
                in_ul = False
            if in_ol:
                html_out.append('</ol>')
                in_ol = False
            if txt.strip():
                html_out.append(f'<p>{txt}</p>')
    if in_ul:
        html_out.append('</ul>')
    if in_ol:
        html_out.append('</ol>')
    # Imagens embutidas
    rels = doc.part.rels
    for rel in rels.values():
        if rel.reltype == 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image':
            img_data = rel.target_part.blob
            img_b64 = base64.b64encode(img_data).decode('utf-8')
            html_out.append(f'<p><img src="data:image/png;base64,{img_b64}" style="max-width:300px;"></p>')
    return jsonify({'html': ''.join(html_out)})

@bp.route('/gpt_assist', methods=['POST'])
def gpt_assist():
    import requests
    api_key = os.environ.get('OPENAI_API_KEY')
    if not api_key:
        return jsonify({'error': 'API key não configurada. Defina OPENAI_API_KEY.'}), 400
    data = request.json
    texto = data.get('texto','')
    acao = data.get('acao','')
    if not texto or not acao:
        return jsonify({'error': 'Texto ou ação não enviados.'}), 400
    prompt_map = {
        'resumir': f"Resuma o seguinte texto de forma clara e concisa:\n{texto}",
        'reescrever': f"Reescreva o texto abaixo de forma mais clara e objetiva:\n{texto}",
        'corrigir': f"Corrija ortografia e gramática do texto abaixo:\n{texto}",
        'traduzir': f"Traduza o texto abaixo para o inglês:\n{texto}",
        'expandir': f"Expanda o texto abaixo com mais detalhes:\n{texto}",
        'gerar_carta': f"Gere uma carta formal ou informal conforme o contexto: {texto}",
        'gerar_email': f"Gere um e-mail profissional ou pessoal conforme o contexto: {texto}",
        'gerar_relatorio': f"Gere um relatório detalhado sobre: {texto}",
        'gerar_roteiro': f"Gere um roteiro (de vídeo, reunião, apresentação, etc.) sobre: {texto}",
        'gerar_lista': f"Gere uma lista (de compras, tarefas, tópicos, etc.) sobre: {texto}",
        'gerar_tabela': f"Gere uma tabela em HTML, com colunas e linhas relevantes, sobre: {texto}",
        'gerar_perguntas': f"Gere perguntas e respostas sobre: {texto}",
        'gerar_livre': f"Gere um texto criativo e relevante sobre: {texto}"
    }
    prompt = prompt_map.get(acao, f"{acao}: {texto}")
    payload = {
        "model": "gpt-3.5-turbo",
        "messages": [
            {"role": "system", "content": "Você é um assistente de texto brasileiro, útil, direto e criativo."},
            {"role": "user", "content": prompt}
        ],
        "max_tokens": 800,
        "temperature": 0.7
    }
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    r = requests.post("https://api.openai.com/v1/chat/completions", json=payload, headers=headers)
    if r.status_code != 200:
        return jsonify({'error': 'Erro na API OpenAI', 'details': r.text}), 500
    resposta = r.json()['choices'][0]['message']['content']
    return jsonify({'resposta': resposta.strip()})

@bp.route('/export_docx', methods=['POST'])
def export_docx():
    html = request.form.get('html', '')
    doc = Document()
    soup = BeautifulSoup(html, 'html.parser')

    def parse_style(style_str):
        # Retorna dict com propriedades CSS relevantes
        style = {}
        if not style_str:
            return style
        for part in style_str.split(';'):
            if ':' in part:
                k, v = part.split(':', 1)
                style[k.strip()] = v.strip()
        return style

    def px_to_pt(px):
        try:
            return int(round(float(px) * 0.75))
        except Exception:
            return None

    def add_run_with_formatting(paragraph, node):
        if not hasattr(node, 'contents') or not node.contents:
            text = node.string or ''
            parts = text.split('\n')
            for i, part in enumerate(parts):
                try:
                    run = paragraph.add_run(part)
                    parent = node.parent if node.parent else node
                    tag = parent.name if hasattr(parent, 'name') else ''
                    style = parse_style(parent.get('style','')) if hasattr(parent, 'get') else {}
                    if tag in ['b', 'strong']:
                        run.bold = True
                    if tag in ['i', 'em']:
                        run.italic = True
                    if tag in ['u']:
                        run.underline = True
                    fs = style.get('font-size')
                    if fs and 'px' in fs:
                        try:
                            sz = px_to_pt(fs.replace('px',''))
                            if sz:
                                run.font.size = Pt(sz)
                        except Exception:
                            pass
                    color = style.get('color')
                    if color and color.startswith('rgb'):
                        m = re.match(r'rgb\((\d+),\s*(\d+),\s*(\d+)\)', color)
                        if m:
                            try:
                                r,g,b = map(int, m.groups())
                                run.font.color.rgb = RGBColor(r,g,b)
                            except Exception:
                                pass
                    if i < len(parts) - 1:
                        run.add_break()
                except Exception:
                    pass
            return
        for child in node.contents:
            if isinstance(child, str):
                parts = child.split('\n')
                for i, part in enumerate(parts):
                    try:
                        run = paragraph.add_run(part)
                        tag = node.name
                        style = parse_style(node.get('style','')) if hasattr(node, 'get') else {}
                        if tag in ['b', 'strong']:
                            run.bold = True
                        if tag in ['i', 'em']:
                            run.italic = True
                        if tag in ['u']:
                            run.underline = True
                        fs = style.get('font-size')
                        if fs and 'px' in fs:
                            try:
                                sz = px_to_pt(fs.replace('px',''))
                                if sz:
                                    run.font.size = Pt(sz)
                            except Exception:
                                pass
                        color = style.get('color')
                        if color and color.startswith('rgb'):
                            m = re.match(r'rgb\((\d+),\s*(\d+),\s*(\d+)\)', color)
                            if m:
                                try:
                                    r,g,b = map(int, m.groups())
                                    run.font.color.rgb = RGBColor(r,g,b)
                                except Exception:
                                    pass
                        if i < len(parts) - 1:
                            run.add_break()
                    except Exception:
                        pass
            elif child.name == 'br':
                try:
                    paragraph.add_run().add_break()
                except Exception:
                    pass
            elif child.name == 'img':
                src = child.get('src', '')
                try:
                    if src.startswith('data:image/'):
                        header, b64data = src.split(',', 1)
                        img_bytes = base64.b64decode(b64data)
                        img_stream = io.BytesIO(img_bytes)
                        # Debug: salvar imagem temporária
                        # with open('debug_img.png', 'wb') as f: f.write(img_bytes)
                        paragraph.add_run().add_picture(img_stream, width=Inches(4))
                    elif src.startswith('http'):
                        resp = requests.get(src)
                        if resp.status_code == 200:
                            img_stream = io.BytesIO(resp.content)
                            # with open('debug_img_url.png', 'wb') as f: f.write(resp.content)
                            paragraph.add_run().add_picture(img_stream, width=Inches(4))
                        else:
                            paragraph.add_run('[Imagem não pôde ser baixada]')
                except Exception as e:
                    paragraph.add_run('[Imagem não pôde ser processada]')
                    print('Erro ao processar imagem:', e)
            else:
                add_run_with_formatting(paragraph, child)

    def process_element(elem, parent=None):
        if elem.name in ['h1','h2','h3','h4','h5','h6']:
            level = int(elem.name[1])
            # Heading com estilo
            style = parse_style(elem.get('style',''))
            p = doc.add_heading('', level=level)
            add_run_with_formatting(p, elem)
            # Alinhamento
            align = style.get('text-align') or elem.get('align')
            if align:
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                if align == 'center':
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif align == 'right':
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif align == 'justify':
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif elem.name in ['ul', 'ol']:
            for li in elem.find_all('li', recursive=False):
                p = doc.add_paragraph('', style='ListBullet' if elem.name=='ul' else 'ListNumber')
                add_run_with_formatting(p, li)
        elif elem.name == 'p' or elem.name == 'div':
            style = parse_style(elem.get('style',''))
            align = elem.get('align') or style.get('text-align')
            p = doc.add_paragraph('')
            if align:
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                if align == 'center':
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif align == 'right':
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif align == 'justify':
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_run_with_formatting(p, elem)
        elif elem.name == 'img':
            # Imagem sozinha: inserir em parágrafo próprio
            src = elem.get('src', '')
            try:
                if src.startswith('data:image/'):
                    header, b64data = src.split(',', 1)
                    img_bytes = base64.b64decode(b64data)
                    img_stream = io.BytesIO(img_bytes)
                    # with open('debug_img_alone.png', 'wb') as f: f.write(img_bytes)
                    doc.add_picture(img_stream, width=Inches(4))
                elif src.startswith('http'):
                    resp = requests.get(src)
                    if resp.status_code == 200:
                        img_stream = io.BytesIO(resp.content)
                        # with open('debug_img_url_alone.png', 'wb') as f: f.write(resp.content)
                        doc.add_picture(img_stream, width=Inches(4))
                    else:
                        doc.add_paragraph('[Imagem não pôde ser baixada]')
            except Exception as e:
                doc.add_paragraph('[Imagem não pôde ser processada]')
                print('Erro ao processar imagem (sozinha):', e)
        elif elem.name:
            p = doc.add_paragraph('')
            add_run_with_formatting(p, elem)
        elif elem.string and elem.string.strip():
            lines = elem.string.split('\n')
            for i, line in enumerate(lines):
                p = doc.add_paragraph(line)
                if i < len(lines) - 1:
                    p.add_run().add_break()

    if soup.body:
        for elem in soup.body.children:
            if getattr(elem, 'name', None) or (getattr(elem, 'string', None) and elem.string.strip()):
                process_element(elem)

    fake_file = io.BytesIO()
    doc.save(fake_file)
    fake_file.seek(0)
    return send_file(fake_file, as_attachment=True, download_name='documento.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@bp.route('/', methods=['GET', 'POST'])
def index():
    plain_text = ''
    html_text = ''
    if request.method == 'POST':
        plain_text = request.form.get('editor', '')
        command = request.form.get('command', '')
        color = request.form.get('color', '')
        align = request.form.get('align', '')
        selection_start = request.form.get('selection_start', None)
        selection_end = request.form.get('selection_end', None)

        html_text = apply_formatting(
            plain_text, command, color, align, selection_start, selection_end
        )
    else:
        preview_body = plain_text.replace('\n', '<br>')
        html_text = f"<html><body>{preview_body}</body></html>"

    return render_template('index.html', content=plain_text, preview=html_text)

def apply_formatting(text, command, color, align, selection_start, selection_end):
    try:
        start = int(selection_start) if selection_start else None
        end = int(selection_end) if selection_end else None
    except Exception:
        start = end = None

    html = text
    if start is not None and end is not None and start < end:
        selected = text[start:end]

        if command == 'bold':
            selected = f'<b>{selected}</b>'
        elif command == 'italic':
            selected = f'<i>{selected}</i>'
        elif command == 'underline':
            selected = f'<u>{selected}</u>'

        if color:
            selected = f'<font color="{color}">{selected}</font>'

        html = text[:start] + selected + text[end:]

    if align:
        html = f'<div align="{align}">{html}</div>'

    html = html.replace('\n', '<br>')
    return f"<html><body>{html}</body></html>"